from flask import Flask, render_template, request, jsonify, g, send_file, redirect, url_for
import os
import re
import pandas as pd
import PyPDF2
from docx import Document
import tempfile
import sqlite3
from datetime import datetime, timedelta
import json
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import csv
from io import StringIO
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps
import secrets

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['DATABASE'] = 'talkkindly.db'
app.config['SECRET_KEY'] = 'your-secret-key-change-this-in-production-12345'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login_page'

class User(UserMixin):
    def __init__(self, id, username, email, is_admin, created_at):
        self.id = id
        self.username = username
        self.email = email
        self.is_admin = bool(is_admin)
        self.created_at = created_at

@login_manager.user_loader
def load_user(user_id):
    db = get_db()
    user = db.execute('SELECT * FROM users WHERE id = ?', (user_id,)).fetchone()
    if user:
        return User(user['id'], user['username'], user['email'], user['is_admin'], user['created_at'])
    return None

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            return jsonify({"error": "Admin access required"}), 403
        return f(*args, **kwargs)
    return decorated_function

def get_db():
    if 'db' not in g:
        g.db = sqlite3.connect(app.config['DATABASE'])
        g.db.row_factory = sqlite3.Row
    return g.db

@app.teardown_appcontext
def close_db(error):
    if hasattr(g, 'db'):
        g.db.close()

def init_db():
    with app.app_context():
        db = get_db()

        db.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                email TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                is_admin BOOLEAN DEFAULT 0,
                is_active BOOLEAN DEFAULT 1,
                created_at TEXT NOT NULL,
                last_login TEXT
            )
        ''')

        db.execute('''
            CREATE TABLE IF NOT EXISTS analyses (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                timestamp TEXT NOT NULL,
                analysis_type TEXT NOT NULL,
                text_content TEXT,
                filename TEXT,
                is_offensive BOOLEAN NOT NULL,
                offensive_count INTEGER NOT NULL,
                details TEXT,
                confidence REAL DEFAULT 0.0,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        ''')

        db.execute('''
            CREATE TABLE IF NOT EXISTS offensive_words (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                word TEXT UNIQUE NOT NULL,
                category TEXT DEFAULT 'general',
                added_by INTEGER,
                added_at TEXT NOT NULL,
                FOREIGN KEY (added_by) REFERENCES users (id)
            )
        ''')

        db.execute('''
            CREATE TABLE IF NOT EXISTS api_keys (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                key TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                created_at TEXT NOT NULL,
                last_used TEXT,
                is_active BOOLEAN DEFAULT 1,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        ''')

        db.execute('''
            CREATE TABLE IF NOT EXISTS activity_log (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                action TEXT NOT NULL,
                details TEXT,
                timestamp TEXT NOT NULL,
                ip_address TEXT,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        ''')

        count = db.execute('SELECT COUNT(*) FROM offensive_words').fetchone()[0]
        if count == 0:
            default_words = [
                ("idiot", "mild"), ("stupid", "mild"), ("dumb", "mild"), ("moron", "mild"),
                ("retard", "mild"), ("fool", "mild"), ("shit", "moderate"), ("ass", "moderate"),
                ("asshole", "moderate"), ("bastard", "moderate"), ("bitch", "moderate"),
                ("damn", "moderate"), ("hell", "moderate"), ("loser", "mild"), ("hate", "mild"),
                ("kill", "moderate"), ("ugly", "mild"), ("fat", "mild"), ("worthless", "mild"),
                ("crap", "mild"), ("fuck", "strong"), ("fucking", "strong"), ("motherfucker", "strong"),
                ("dick", "strong"), ("pussy", "strong"), ("cock", "strong"), ("whore", "strong"),
                ("slut", "strong"), ("douchebag", "moderate"), ("scumbag", "moderate"),
                ("cunt", "strong")
            ]
            for word, category in default_words:
                db.execute(
                    'INSERT OR IGNORE INTO offensive_words (word, category, added_at) VALUES (?, ?, ?)',
                    (word, category, datetime.now().isoformat())
                )

        db.commit()
        print("✅ Database ready")

def get_offensive_words():
    db = get_db()
    words = db.execute('SELECT word FROM offensive_words').fetchall()
    return set([word['word'] for word in words])

def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r'[^\w\s]', ' ', text)
    text = re.sub(r'\d+', '', text)
    text = ' '.join(text.split())
    return text

def detect_with_patterns(text: str) -> dict:
    if not text or not text.strip():
        return {"found": [], "patterns_found": [], "suspicious_words": [], "count": 0, "is_offensive": False}

    OFFENSIVE_WORDS = get_offensive_words()
    cleaned = clean_text(text)
    words = cleaned.split()
    
    found_words = [word for word in words if word in OFFENSIVE_WORDS]
    
    found_patterns = []
    flexible_patterns = [
        r'\b[a-z]{2,}[\*]+[a-z]*[\*]*[a-z]*\b',
        r'\b[a-z]*[\$\!]+[a-z]+\b',
        r'\b[a-z]+\d+[a-z]*\b',
        r'\b[a-z]*[\*\$\!\d]+[a-z]*[\*\$\!\d]*[a-z]*\b'
    ]
    
    for pattern in flexible_patterns:
        matches = re.findall(pattern, text.lower())
        for match in matches:
            if len(re.sub(r'[\*\$\!\d]', '', match)) >= 2:
                found_patterns.append(match)
    
    suspicious_words = []
    offensive_roots = list(OFFENSIVE_WORDS)
    
    for word in words:
        for root in offensive_roots:
            if root in word and len(word) >= len(root) - 1:
                if word not in found_words and word not in suspicious_words:
                    suspicious_words.append(word)
                break
    
    all_offensive = found_words + found_patterns + suspicious_words
    
    return {
        "found": found_words,
        "patterns_found": found_patterns,
        "suspicious_words": suspicious_words,
        "count": len(all_offensive),
        "is_offensive": len(all_offensive) > 0,
        "clean_text": cleaned
    }

def calculate_confidence(result: dict) -> float:
    base_confidence = 0.0
    if result['found']:
        base_confidence += 0.6
    if result['patterns_found']:
        base_confidence += 0.3
    if result['suspicious_words']:
        base_confidence += 0.1
    return min(base_confidence, 1.0)

def log_activity(user_id, action, details=None, ip_address=None):
    try:
        db = get_db()
        db.execute('''
            INSERT INTO activity_log (user_id, action, details, timestamp, ip_address)
            VALUES (?, ?, ?, ?, ?)
        ''', (user_id, action, json.dumps(details) if details else None, 
              datetime.now().isoformat(), ip_address))
        db.commit()
    except Exception as e:
        print(f"Error logging activity: {e}")

def save_analysis(analysis_data):
    db = get_db()
    cursor = db.execute('''
        INSERT INTO analyses 
        (user_id, timestamp, analysis_type, text_content, filename, is_offensive, offensive_count, details, confidence)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        current_user.id,
        datetime.now().isoformat(),
        analysis_data['analysis_type'],
        analysis_data.get('text_content'),
        analysis_data.get('filename'),
        analysis_data['is_offensive'],
        analysis_data['offensive_count'],
        json.dumps(analysis_data.get('details', {})),
        analysis_data.get('confidence', 0.0)
    ))
    db.commit()
    
    log_activity(current_user.id, 'analysis', {
        'type': analysis_data['analysis_type'],
        'is_offensive': analysis_data['is_offensive'],
        'offensive_count': analysis_data['offensive_count']
    }, request.remote_addr)
    
    return cursor.lastrowid

def get_recent_analyses(user_id=None, limit=50, filter_type=None, search_query=None, start_date=None, end_date=None):
    db = get_db()
    
    if user_id and not current_user.is_admin:
        query = '''
            SELECT a.*, u.username 
            FROM analyses a 
            JOIN users u ON a.user_id = u.id 
            WHERE a.user_id = ?
        '''
        params = [user_id]
    else:
        query = 'SELECT a.*, u.username FROM analyses a JOIN users u ON a.user_id = u.id WHERE 1=1'
        params = []
    
    if filter_type and filter_type != 'all':
        query += ' AND a.analysis_type = ?'
        params.append(filter_type)
    
    if search_query:
        query += ' AND (a.text_content LIKE ? OR a.filename LIKE ? OR u.username LIKE ?)'
        params.extend([f'%{search_query}%', f'%{search_query}%', f'%{search_query}%'])
    
    if start_date:
        query += ' AND DATE(a.timestamp) >= ?'
        params.append(start_date)
    if end_date:
        query += ' AND DATE(a.timestamp) <= ?'
        params.append(end_date)
    
    query += ' ORDER BY a.timestamp DESC LIMIT ?'
    params.append(limit)
    
    analyses = db.execute(query, params).fetchall()
    return [dict(analysis) for analysis in analyses]

def get_statistics(user_id=None):
    db = get_db()
    
    if user_id:
        query = '''
            SELECT 
                COUNT(*) as total_analyses,
                SUM(CASE WHEN is_offensive = 1 THEN 1 ELSE 0 END) as offensive_analyses,
                SUM(CASE WHEN is_offensive = 0 THEN 1 ELSE 0 END) as safe_analyses
            FROM analyses
            WHERE user_id = ?
        '''
        params = [user_id]
    else:
        query = '''
            SELECT 
                COUNT(*) as total_analyses,
                SUM(CASE WHEN is_offensive = 1 THEN 1 ELSE 0 END) as offensive_analyses,
                SUM(CASE WHEN is_offensive = 0 THEN 1 ELSE 0 END) as safe_analyses
            FROM analyses
        '''
        params = []
    
    stats = db.execute(query, params).fetchone()
    return dict(stats) if stats else {"total_analyses": 0, "offensive_analyses": 0, "safe_analyses": 0}

def get_detailed_statistics(user_id=None):
    overall = get_statistics(user_id)
    
    db = get_db()
    
    if user_id:
        query = 'SELECT details FROM analyses WHERE is_offensive = 1 AND user_id = ?'
        params = [user_id]
    else:
        query = 'SELECT details FROM analyses WHERE is_offensive = 1'
        params = []
    
    analyses = db.execute(query, params).fetchall()
    
    direct_count = 0
    pattern_count = 0
    suspicious_count = 0
    
    for analysis in analyses:
        try:
            details = json.loads(analysis['details'])
            if details.get('words'):
                direct_count += len(details['words'])
            if details.get('patterns'):
                pattern_count += len(details['patterns'])
            if details.get('suspicious'):
                suspicious_count += len(details['suspicious'])
        except:
            continue
    
    return {
        **overall,
        'direct_words': direct_count,
        'hidden_patterns': pattern_count,
        'suspicious_words': suspicious_count
    }

def process_csv_file(file_path: str) -> dict:
    try:
        df = pd.read_csv(file_path)
        results = []
        total_offensive = 0
        all_offensive_details = []
        
        for index, row in df.iterrows():
            text = ""
            if len(row) > 0:
                for col in row.index:
                    if isinstance(row[col], str) and len(row[col].strip()) > 0:
                        text = str(row[col])
                        break
            
            if text:
                analysis = detect_with_patterns(text)
                results.append({
                    "row": index + 1,
                    "text": text[:100] + "..." if len(text) > 100 else text,
                    "analysis": analysis
                })
                if analysis["is_offensive"]:
                    total_offensive += 1
                    if analysis['found']:
                        all_offensive_details.extend(analysis['found'])
        
        return {
            "file_type": "CSV",
            "total_rows": len(df),
            "offensive_rows": total_offensive,
            "offensive_words": list(set(all_offensive_details)),
            "results": results,
            "safe_percentage": ((len(df) - total_offensive) / len(df)) * 100 if len(df) > 0 else 100,
            "analysis": {
                "is_offensive": total_offensive > 0,
                "count": total_offensive,
                "found": all_offensive_details
            }
        }
    except Exception as e:
        return {"error": f"CSV processing error: {str(e)}"}

def process_pdf_file(file_path: str) -> dict:
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            analysis = detect_with_patterns(text)
            
            return {
                "file_type": "PDF",
                "total_pages": len(pdf_reader.pages),
                "analysis": analysis,
                "text_preview": text[:200] + "..." if len(text) > 200 else text
            }
    except Exception as e:
        return {"error": f"PDF processing error: {str(e)}"}

def process_docx_file(file_path: str) -> dict:
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        analysis = detect_with_patterns(text)
        
        return {
            "file_type": "DOCX",
            "total_paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
            "analysis": analysis,
            "text_preview": text[:200] + "..." if len(text) > 200 else text
        }
    except Exception as e:
        return {"error": f"DOCX processing error: {str(e)}"}

def process_text_file(file_path: str) -> dict:
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            text = file.read()
        
        analysis = detect_with_patterns(text)
        
        return {
            "file_type": "TXT",
            "analysis": analysis,
            "text_preview": text[:200] + "..." if len(text) > 200 else text
        }
    except Exception as e:
        return {"error": f"Text file processing error: {str(e)}"}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'txt', 'csv', 'pdf', 'docx'}

# ==================== ROUTES ====================

@app.route('/')
def home():
    if not current_user.is_authenticated:
        return redirect(url_for('login_page'))
    return render_template('index.html', user=current_user)

@app.route('/login')
def login_page():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    return render_template('login.html')

@app.route('/register')
def register_page():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    return render_template('register.html')

@app.route('/admin')
@login_required
@admin_required
def admin_panel():
    return render_template('admin.html', user=current_user)

# ==================== AUTH API ====================

@app.route('/api/register', methods=['POST'])
def register():
    try:
        data = request.get_json()
        
        username = data.get('username', '').strip()
        email = data.get('email', '').strip()
        password = data.get('password', '')
        
        if not username or not email or not password:
            return jsonify({"error": "All fields are required"}), 400
        
        if len(username) < 3:
            return jsonify({"error": "Username must be at least 3 characters"}), 400
        
        if len(password) < 6:
            return jsonify({"error": "Password must be at least 6 characters"}), 400
        
        if '@' not in email:
            return jsonify({"error": "Invalid email address"}), 400
        
        db = get_db()
        
        existing_user = db.execute('SELECT id FROM users WHERE username = ? OR email = ?', 
                                   (username, email)).fetchone()
        
        if existing_user:
            return jsonify({"error": "Username or email already exists"}), 400
        
        user_count = db.execute('SELECT COUNT(*) as count FROM users').fetchone()['count']
        is_admin = (user_count == 0)
        
        print(f"👑 User registration - Total users: {user_count}, Is Admin: {is_admin}")
        
        password_hash = generate_password_hash(password)
        
        cursor = db.execute('''
            INSERT INTO users (username, email, password_hash, is_admin, created_at)
            VALUES (?, ?, ?, ?, ?)
        ''', (username, email, password_hash, is_admin, datetime.now().isoformat()))
        db.commit()
        
        user_id = cursor.lastrowid
        user = User(user_id, username, email, is_admin, datetime.now().isoformat())
        login_user(user, remember=True)
        
        log_activity(user_id, 'register', {'email': email, 'is_admin': is_admin}, request.remote_addr)
        
        return jsonify({
            "success": True,
            "message": "Registration successful!" + (" You are the admin!" if is_admin else ""),
            "user": {
                "id": user_id,
                "username": username,
                "email": email,
                "is_admin": is_admin
            }
        })
    except Exception as e:
        print(f"❌ Registration error: {e}")
        return jsonify({"error": f"Registration failed: {str(e)}"}), 500

@app.route('/api/login', methods=['POST'])
def login():
    try:
        data = request.get_json()
        
        username = data.get('username', '').strip()
        password = data.get('password', '')
        remember = data.get('remember', False)
        
        if not username or not password:
            return jsonify({"error": "Username and password are required"}), 400
        
        db = get_db()
        user_data = db.execute('SELECT * FROM users WHERE username = ? OR email = ?', 
                               (username, username)).fetchone()
        
        if not user_data:
            return jsonify({"error": "Invalid username or password"}), 401
        
        if not user_data['is_active']:
            return jsonify({"error": "Account is disabled"}), 403
        
        if not check_password_hash(user_data['password_hash'], password):
            return jsonify({"error": "Invalid username or password"}), 401
        
        db.execute('UPDATE users SET last_login = ? WHERE id = ?', 
                   (datetime.now().isoformat(), user_data['id']))
        db.commit()
        
        user = User(user_data['id'], user_data['username'], user_data['email'], 
                    user_data['is_admin'], user_data['created_at'])
        login_user(user, remember=remember)
        
        log_activity(user.id, 'login', {'is_admin': user.is_admin}, request.remote_addr)
        
        return jsonify({
            "success": True,
            "message": "Login successful!",
            "user": {
                "id": user.id,
                "username": user.username,
                "email": user.email,
                "is_admin": user.is_admin
            }
        })
    except Exception as e:
        print(f"❌ Login error: {e}")
        return jsonify({"error": f"Login failed: {str(e)}"}), 500

@app.route('/api/logout', methods=['POST'])
@login_required
def logout():
    log_activity(current_user.id, 'logout', None, request.remote_addr)
    logout_user()
    return jsonify({"success": True, "message": "Logged out successfully"})

@app.route('/api/user/profile', methods=['GET'])
@login_required
def get_profile():
    try:
        db = get_db()
        user_data = db.execute('SELECT * FROM users WHERE id = ?', (current_user.id,)).fetchone()
        
        if not user_data:
            return jsonify({"error": "User not found"}), 404
        
        stats = get_statistics(current_user.id)
        
        return jsonify({
            "id": user_data['id'],
            "username": user_data['username'],
            "email": user_data['email'],
            "is_admin": bool(user_data['is_admin']),
            "created_at": user_data['created_at'],
            "last_login": user_data['last_login'],
            "statistics": stats
        })
    except Exception as e:
        print(f"❌ Profile error: {e}")
        return jsonify({"error": f"Error loading profile: {str(e)}"}), 500

# ==================== ANALYSIS API ====================

@app.route('/analyze/text', methods=['POST'])
@login_required
def analyze_text():
    try:
        text = ""
        
        if request.content_type and 'application/json' in request.content_type:
            data = request.get_json()
            text = data.get('text', '')
        else:
            text = request.form.get('text', '')
        
        if not text:
            return jsonify({"error": "No text provided"}), 400
        
        print(f"📨 Analyzing text: {text[:50]}...")
        
        result = detect_with_patterns(text)
        confidence = calculate_confidence(result)
        
        analysis_id = save_analysis({
            'analysis_type': 'text',
            'text_content': text,
            'is_offensive': result['is_offensive'],
            'offensive_count': result['count'],
            'details': {
                'words': result['found'],
                'patterns': result['patterns_found'],
                'suspicious': result['suspicious_words']
            },
            'confidence': confidence
        })
        
        if result['is_offensive']:
            message = f"🚫 Offensive content detected! Found {result['count']} offensive items."
            details = []
            if result['found']:
                details.append(f"Explicit words: {', '.join(result['found'])}")
            if result['patterns_found']:
                details.append(f"Hidden patterns: {', '.join(result['patterns_found'])}")
            if result['suspicious_words']:
                details.append(f"Suspicious words: {', '.join(result['suspicious_words'])}")
            
            if details:
                message += " Details: " + "; ".join(details)
        else:
            message = "✅ Text is safe and contains no offensive content"
        
        print(f"✅ Analysis completed: {message}")
        
        return jsonify({
            "result": message,
            "is_offensive": result['is_offensive'],
            "offensive_words": result['found'],
            "offensive_patterns": result['patterns_found'],
            "suspicious_words": result['suspicious_words'],
            "word_count": result['count'],
            "confidence": confidence,
            "analysis_id": analysis_id
        })
        
    except Exception as e:
        print(f"❌ Analysis error: {e}")
        return jsonify({"error": f"Analysis error: {str(e)}"}), 500

@app.route('/analyze/file', methods=['POST'])
@login_required
def analyze_file():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        if file and allowed_file(file.filename):
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(file_path)
            
            file_extension = file.filename.rsplit('.', 1)[1].lower()
            print(f"📁 Processing {file_extension} file: {file.filename}")
            
            if file_extension == 'csv':
                result = process_csv_file(file_path)
            elif file_extension == 'pdf':
                result = process_pdf_file(file_path)
            elif file_extension == 'docx':
                result = process_docx_file(file_path)
            else:
                result = process_text_file(file_path)
            
            if 'error' not in result:
                analysis_data = result.get('analysis', {})
                save_analysis({
                    'analysis_type': 'file',
                    'filename': file.filename,
                    'is_offensive': analysis_data.get('is_offensive', False),
                    'offensive_count': analysis_data.get('count', 0),
                    'details': analysis_data
                })
            
            try:
                os.remove(file_path)
            except:
                pass
            
            return jsonify(result)
        
        return jsonify({"error": "File type not allowed. Use CSV, PDF, DOCX, or TXT"}), 400
    except Exception as e:
        print(f"❌ File analysis error: {e}")
        return jsonify({"error": str(e)}), 500

# ==================== HISTORY & STATS API ====================

@app.route('/api/history', methods=['GET'])
@login_required
def api_history():
    try:
        filter_type = request.args.get('type', 'all')
        search_query = request.args.get('search', None)
        start_date = request.args.get('start_date', None)
        end_date = request.args.get('end_date', None)
        limit = int(request.args.get('limit', 50))
        
        user_id = None if current_user.is_admin else current_user.id
        
        analyses = get_recent_analyses(
            user_id=user_id,
            limit=limit,
            filter_type=filter_type,
            search_query=search_query,
            start_date=start_date,
            end_date=end_date
        )
        return jsonify(analyses)
    except Exception as e:
        print(f"❌ History error: {e}")
        return jsonify({"error": f"Error loading history: {str(e)}"}), 500

@app.route('/api/statistics', methods=['GET'])
@login_required
def api_statistics():
    try:
        user_id = None if current_user.is_admin else current_user.id
        stats = get_detailed_statistics(user_id)
        return jsonify(stats)
    except Exception as e:
        print(f"❌ Statistics error: {e}")
        return jsonify({"error": f"Error loading statistics: {str(e)}"}), 500

@app.route('/api/insights', methods=['GET'])
@login_required
def api_insights():
    try:
        db = get_db()
        
        if current_user.is_admin:
            analyses = db.execute('SELECT details FROM analyses WHERE is_offensive = 1').fetchall()
        else:
            analyses = db.execute('''
                SELECT details FROM analyses 
                WHERE is_offensive = 1 AND user_id = ?
            ''', (current_user.id,)).fetchall()
        
        word_count = {}
        for analysis in analyses:
            try:
                details = json.loads(analysis['details'])
                for word in details.get('words', []):
                    word_count[word] = word_count.get(word, 0) + 1
            except:
                continue
        
        top_words = sorted(word_count.items(), key=lambda x: x[1], reverse=True)[:10]
        
        user_id = None if current_user.is_admin else current_user.id
        
        return jsonify({
            "top_offensive_words": top_words,
            "statistics": get_detailed_statistics(user_id)
        })
    except Exception as e:
        print(f"❌ Insights error: {e}")
        return jsonify({"error": f"Error loading insights: {str(e)}"}), 500

@app.route('/api/reanalyze/<int:analysis_id>', methods=['POST'])
@login_required
def reanalyze(analysis_id):
    try:
        db = get_db()
        
        if current_user.is_admin:
            analysis = db.execute('SELECT * FROM analyses WHERE id = ?', (analysis_id,)).fetchone()
        else:
            analysis = db.execute('SELECT * FROM analyses WHERE id = ? AND user_id = ?', 
                                 (analysis_id, current_user.id)).fetchone()
        
        if not analysis:
            return jsonify({"error": "Analysis not found"}), 404
        
        text = analysis['text_content']
        if not text:
            return jsonify({"error": "No text content to re-analyze"}), 400
        
        result = detect_with_patterns(text)
        confidence = calculate_confidence(result)
        
        new_id = save_analysis({
            'analysis_type': 'text',
            'text_content': text,
            'is_offensive': result['is_offensive'],
            'offensive_count': result['count'],
            'details': {
                'words': result['found'],
                'patterns': result['patterns_found'],
                'suspicious': result['suspicious_words']
            },
            'confidence': confidence
        })
        
        return jsonify({
            "success": True,
            "new_analysis_id": new_id,
            "result": result,
            "confidence": confidence
        })
    except Exception as e:
        print(f"❌ Reanalyze error: {e}")
        return jsonify({"error": str(e)}), 500

# ==================== ADMIN API ====================

@app.route('/api/admin/statistics', methods=['GET'])
@login_required
@admin_required
def get_admin_statistics():
    try:
        db = get_db()
        
        # System statistics
        stats = db.execute('''
            SELECT 
                COUNT(*) as total_analyses,
                SUM(CASE WHEN is_offensive = 1 THEN 1 ELSE 0 END) as offensive_analyses,
                SUM(CASE WHEN is_offensive = 0 THEN 1 ELSE 0 END) as safe_analyses
            FROM analyses
        ''').fetchone()
        
        # User statistics
        user_stats = db.execute('''
            SELECT 
                COUNT(*) as total_users,
                SUM(CASE WHEN is_admin = 1 THEN 1 ELSE 0 END) as admin_count,
                SUM(CASE WHEN is_active = 1 THEN 1 ELSE 0 END) as active_users
            FROM users
        ''').fetchone()
        
        # Today's analyses
        today = datetime.now().date().isoformat()
        today_stats = db.execute('''
            SELECT COUNT(*) as today_analyses
            FROM analyses
            WHERE DATE(timestamp) = ?
        ''', (today,)).fetchone()
        
        # Offensive words count
        words_count = db.execute('SELECT COUNT(*) as count FROM offensive_words').fetchone()
        
        # Detailed statistics
        detailed_stats = get_detailed_statistics()
        
        return jsonify({
            'total_users': user_stats['total_users'] if user_stats else 0,
            'total_analyses': stats['total_analyses'] if stats else 0,
            'offensive_analyses': stats['offensive_analyses'] if stats else 0,
            'safe_analyses': stats['safe_analyses'] if stats else 0,
            'today_analyses': today_stats['today_analyses'] if today_stats else 0,
            'offensive_words_count': words_count['count'] if words_count else 0,
            'admin_count': user_stats['admin_count'] if user_stats else 0,
            'active_users': user_stats['active_users'] if user_stats else 0,
            'direct_words': detailed_stats.get('direct_words', 0),
            'hidden_patterns': detailed_stats.get('hidden_patterns', 0),
            'suspicious_words': detailed_stats.get('suspicious_words', 0)
        })
    except Exception as e:
        print(f"❌ Admin statistics error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/users', methods=['GET'])
@login_required
@admin_required
def get_all_users():
    try:
        db = get_db()
        users = db.execute('''
            SELECT u.*, 
                   COUNT(a.id) as total_analyses,
                   SUM(CASE WHEN a.is_offensive = 1 THEN 1 ELSE 0 END) as offensive_analyses
            FROM users u
            LEFT JOIN analyses a ON u.id = a.user_id
            GROUP BY u.id
            ORDER BY u.created_at DESC
        ''').fetchall()
        
        users_list = []
        for user in users:
            user_dict = dict(user)
            user_dict['is_admin'] = bool(user['is_admin'])
            user_dict['is_active'] = bool(user['is_active'])
            users_list.append(user_dict)
            
        return jsonify(users_list)
    except Exception as e:
        print(f"❌ Get users error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/users/<int:user_id>/toggle', methods=['POST'])
@login_required
@admin_required
def toggle_user_status(user_id):
    if user_id == current_user.id:
        return jsonify({"error": "Cannot disable your own account"}), 400
    
    try:
        db = get_db()
        user = db.execute('SELECT * FROM users WHERE id = ?', (user_id,)).fetchone()
        
        if not user:
            return jsonify({"error": "User not found"}), 404
        
        new_status = not user['is_active']
        db.execute('UPDATE users SET is_active = ? WHERE id = ?', (new_status, user_id))
        db.commit()
        
        log_activity(current_user.id, 'toggle_user', {
            'user_id': user_id,
            'username': user['username'],
            'new_status': 'active' if new_status else 'disabled'
        }, request.remote_addr)
        
        return jsonify({"success": True, "is_active": new_status})
    except Exception as e:
        print(f"❌ Toggle user error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/users/<int:user_id>', methods=['DELETE'])
@login_required
@admin_required
def delete_user(user_id):
    if user_id == current_user.id:
        return jsonify({"error": "Cannot delete your own account"}), 400
    
    try:
        db = get_db()
        user = db.execute('SELECT username FROM users WHERE id = ?', (user_id,)).fetchone()
        
        if not user:
            return jsonify({"error": "User not found"}), 404
        
        db.execute('DELETE FROM analyses WHERE user_id = ?', (user_id,))
        db.execute('DELETE FROM api_keys WHERE user_id = ?', (user_id,))
        db.execute('DELETE FROM activity_log WHERE user_id = ?', (user_id,))
        db.execute('DELETE FROM users WHERE id = ?', (user_id,))
        db.commit()
        
        log_activity(current_user.id, 'delete_user', {
            'user_id': user_id,
            'username': user['username']
        }, request.remote_addr)
        
        return jsonify({"success": True})
    except Exception as e:
        print(f"❌ Delete user error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/words', methods=['GET'])
@login_required
@admin_required
def get_offensive_words_list():
    try:
        db = get_db()
        words = db.execute('''
            SELECT w.*, u.username as added_by_name
            FROM offensive_words w
            LEFT JOIN users u ON w.added_by = u.id
            ORDER BY w.word
        ''').fetchall()
        return jsonify([dict(word) for word in words])
    except Exception as e:
        print(f"❌ Get words error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/words', methods=['POST'])
@login_required
@admin_required
def add_offensive_word():
    try:
        data = request.get_json()
        word = data.get('word', '').strip().lower()
        category = data.get('category', 'general')
        
        if not word:
            return jsonify({"error": "Word is required"}), 400
        
        db = get_db()
        
        existing = db.execute('SELECT id FROM offensive_words WHERE word = ?', (word,)).fetchone()
        if existing:
            return jsonify({"error": "Word already exists"}), 400
        
        db.execute('''
            INSERT INTO offensive_words (word, category, added_by, added_at)
            VALUES (?, ?, ?, ?)
        ''', (word, category, current_user.id, datetime.now().isoformat()))
        db.commit()
        
        log_activity(current_user.id, 'add_word', {'word': word, 'category': category}, request.remote_addr)
        
        return jsonify({"success": True, "message": f"Word '{word}' added successfully"})
    except Exception as e:
        print(f"❌ Add word error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/words/<int:word_id>', methods=['DELETE'])
@login_required
@admin_required
def delete_offensive_word(word_id):
    try:
        db = get_db()
        word = db.execute('SELECT word FROM offensive_words WHERE id = ?', (word_id,)).fetchone()
        
        if not word:
            return jsonify({"error": "Word not found"}), 404
        
        db.execute('DELETE FROM offensive_words WHERE id = ?', (word_id,))
        db.commit()
        
        log_activity(current_user.id, 'delete_word', {'word': word['word']}, request.remote_addr)
        
        return jsonify({"success": True})
    except Exception as e:
        print(f"❌ Delete word error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/analyses', methods=['GET'])
@login_required
@admin_required
def get_all_analyses():
    try:
        db = get_db()
        limit = int(request.args.get('limit', 100))
        search = request.args.get('search', '')
        analysis_type = request.args.get('type', 'all')
        
        query = '''
            SELECT a.*, u.username, u.is_admin
            FROM analyses a
            JOIN users u ON a.user_id = u.id
            WHERE 1=1
        '''
        params = []
        
        if search:
            query += ' AND (a.text_content LIKE ? OR u.username LIKE ? OR a.filename LIKE ?)'
            params.extend([f'%{search}%', f'%{search}%', f'%{search}%'])
        
        if analysis_type != 'all':
            query += ' AND a.analysis_type = ?'
            params.append(analysis_type)
        
        query += ' ORDER BY a.timestamp DESC LIMIT ?'
        params.append(limit)
        
        analyses = db.execute(query, params).fetchall()
        
        result = []
        for analysis in analyses:
            analysis_dict = dict(analysis)
            if analysis_dict.get('details'):
                try:
                    analysis_dict['details'] = json.loads(analysis_dict['details'])
                except:
                    analysis_dict['details'] = {}
            result.append(analysis_dict)
        
        return jsonify(result)
    except Exception as e:
        print(f"❌ Admin analyses error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/analyses/<int:analysis_id>', methods=['DELETE'])
@login_required
@admin_required
def delete_analysis(analysis_id):
    try:
        db = get_db()
        
        analysis = db.execute('SELECT a.*, u.username FROM analyses a JOIN users u ON a.user_id = u.id WHERE a.id = ?', 
                             (analysis_id,)).fetchone()
        
        if not analysis:
            return jsonify({"error": "Analysis not found"}), 404
        
        db.execute('DELETE FROM analyses WHERE id = ?', (analysis_id,))
        db.commit()
        
        log_activity(current_user.id, 'delete_analysis', {
            'analysis_id': analysis_id,
            'username': analysis['username'],
            'text_preview': analysis['text_content'][:50] + '...' if analysis['text_content'] else 'File analysis'
        }, request.remote_addr)
        
        return jsonify({"success": True, "message": "Analysis deleted successfully"})
    except Exception as e:
        print(f"❌ Delete analysis error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/admin/activity', methods=['GET'])
@login_required
@admin_required
def get_activity_log():
    try:
        db = get_db()
        limit = int(request.args.get('limit', 50))
        
        activities = db.execute('''
            SELECT a.*, u.username
            FROM activity_log a
            JOIN users u ON a.user_id = u.id
            ORDER BY a.timestamp DESC
            LIMIT ?
        ''', (limit,)).fetchall()
        
        return jsonify([dict(activity) for activity in activities])
    except Exception as e:
        print(f"❌ Activity log error: {e}")
        return jsonify({"error": str(e)}), 500

# ==================== API KEYS ====================

@app.route('/api/keys', methods=['GET'])
@login_required
def get_api_keys():
    try:
        db = get_db()
        keys = db.execute('''
            SELECT id, name, key, created_at, last_used, is_active
            FROM api_keys
            WHERE user_id = ?
            ORDER BY created_at DESC
        ''', (current_user.id,)).fetchall()
        return jsonify([dict(key) for key in keys])
    except Exception as e:
        print(f"❌ Get API keys error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/keys', methods=['POST'])
@login_required
def create_api_key():
    try:
        data = request.get_json()
        name = data.get('name', '').strip()
        
        if not name:
            return jsonify({"error": "Key name is required"}), 400
        
        api_key = 'tk_' + secrets.token_urlsafe(32)
        
        db = get_db()
        cursor = db.execute('''
            INSERT INTO api_keys (user_id, key, name, created_at, is_active)
            VALUES (?, ?, ?, ?, 1)
        ''', (current_user.id, api_key, name, datetime.now().isoformat()))
        db.commit()
        
        log_activity(current_user.id, 'create_api_key', {'name': name}, request.remote_addr)
        
        return jsonify({
            "success": True,
            "key": api_key,
            "message": "API key created successfully"
        })
    except Exception as e:
        print(f"❌ Create API key error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/keys/<int:key_id>', methods=['DELETE'])
@login_required
def delete_api_key(key_id):
    try:
        db = get_db()
        
        key = db.execute('SELECT name FROM api_keys WHERE id = ? AND user_id = ?', 
                         (key_id, current_user.id)).fetchone()
        
        if not key:
            return jsonify({"error": "Key not found"}), 404
        
        db.execute('DELETE FROM api_keys WHERE id = ? AND user_id = ?', 
                   (key_id, current_user.id))
        db.commit()
        
        log_activity(current_user.id, 'delete_api_key', {'name': key['name']}, request.remote_addr)
        
        return jsonify({"success": True})
    except Exception as e:
        print(f"❌ Delete API key error: {e}")
        return jsonify({"error": str(e)}), 500

# ==================== EXPORT & EXTERNAL API ====================

@app.route('/api/export/pdf', methods=['POST'])
@login_required
def export_pdf():
    try:
        data = request.get_json()
        analysis_ids = data.get('ids', [])
        
        db = get_db()
        
        if analysis_ids:
            placeholders = ','.join('?' * len(analysis_ids))
            if current_user.is_admin:
                analyses = db.execute(f'SELECT * FROM analyses WHERE id IN ({placeholders}) ORDER BY timestamp DESC', 
                                     analysis_ids).fetchall()
            else:
                analyses = db.execute(f'SELECT * FROM analyses WHERE id IN ({placeholders}) AND user_id = ? ORDER BY timestamp DESC', 
                                     (*analysis_ids, current_user.id)).fetchall()
        else:
            if current_user.is_admin:
                analyses = db.execute('SELECT * FROM analyses ORDER BY timestamp DESC LIMIT 50').fetchall()
            else:
                analyses = db.execute('SELECT * FROM analyses WHERE user_id = ? ORDER BY timestamp DESC LIMIT 50', 
                                    (current_user.id,)).fetchall()
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#2c3e50'))
        
        elements.append(Paragraph(f"TalkKindly - Analysis Report", title_style))
        elements.append(Spacer(1, 0.2*inch))
        elements.append(Paragraph(f"User: {current_user.username}", styles['Normal']))
        elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        elements.append(Spacer(1, 0.5*inch))
        
        user_id = None if current_user.is_admin else current_user.id
        stats = get_statistics(user_id)
        summary_data = [
            ['Total Analyses', str(stats['total_analyses'])],
            ['Offensive', str(stats['offensive_analyses'])],
            ['Safe', str(stats['safe_analyses'])]
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#ecf0f1')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#2c3e50')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7'))
        ]))
        
        elements.append(summary_table)
        elements.append(Spacer(1, 0.5*inch))
        
        elements.append(Paragraph("Analysis Details", styles['Heading2']))
        elements.append(Spacer(1, 0.2*inch))
        
        for analysis in analyses:
            details_text = f"<b>ID:</b> {analysis['id']}<br/>"
            details_text += f"<b>Date:</b> {analysis['timestamp']}<br/>"
            details_text += f"<b>Type:</b> {analysis['analysis_type']}<br/>"
            details_text += f"<b>Status:</b> {'Offensive' if analysis['is_offensive'] else 'Safe'}<br/>"
            details_text += f"<b>Count:</b> {analysis['offensive_count']}<br/>"
            
            if analysis['text_content']:
                preview = analysis['text_content'][:100] + "..." if len(analysis['text_content']) > 100 else analysis['text_content']
                details_text += f"<b>Content:</b> {preview}"
            
            elements.append(Paragraph(details_text, styles['Normal']))
            elements.append(Spacer(1, 0.3*inch))
        
        doc.build(elements)
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'talkkindly_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        )
    except Exception as e:
        print(f"❌ PDF export error: {e}")
        return jsonify({"error": f"PDF export error: {str(e)}"}), 500

@app.route('/api/export/csv', methods=['POST'])
@login_required
def export_csv():
    try:
        data = request.get_json()
        analysis_ids = data.get('ids', [])
        
        db = get_db()
        
        if analysis_ids:
            placeholders = ','.join('?' * len(analysis_ids))
            if current_user.is_admin:
                analyses = db.execute(f'SELECT * FROM analyses WHERE id IN ({placeholders}) ORDER BY timestamp DESC', 
                                     analysis_ids).fetchall()
            else:
                analyses = db.execute(f'SELECT * FROM analyses WHERE id IN ({placeholders}) AND user_id = ? ORDER BY timestamp DESC', 
                                     (*analysis_ids, current_user.id)).fetchall()
        else:
            if current_user.is_admin:
                analyses = db.execute('SELECT * FROM analyses ORDER BY timestamp DESC LIMIT 50').fetchall()
            else:
                analyses = db.execute('SELECT * FROM analyses WHERE user_id = ? ORDER BY timestamp DESC LIMIT 50', 
                                    (current_user.id,)).fetchall()
        
        output = StringIO()
        writer = csv.writer(output)
        
        writer.writerow(['ID', 'User ID', 'Timestamp', 'Type', 'Filename', 'Text Content', 'Is Offensive', 'Offensive Count', 'Confidence'])
        
        for analysis in analyses:
            writer.writerow([
                analysis['id'],
                analysis['user_id'],
                analysis['timestamp'],
                analysis['analysis_type'],
                analysis['filename'] or '',
                (analysis['text_content'][:100] + '...') if analysis['text_content'] and len(analysis['text_content']) > 100 else (analysis['text_content'] or ''),
                'Yes' if analysis['is_offensive'] else 'No',
                analysis['offensive_count'],
                f"{analysis['confidence']:.2f}" if analysis['confidence'] else '0.00'
            ])
        
        output.seek(0)
        
        return send_file(
            BytesIO(output.getvalue().encode('utf-8')),
            mimetype='text/csv',
            as_attachment=True,
            download_name=f'talkkindly_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
        )
    except Exception as e:
        print(f"❌ CSV export error: {e}")
        return jsonify({"error": f"CSV export error: {str(e)}"}), 500

@app.route('/api/analyze', methods=['POST'])
def api_analyze():
    try:
        if request.is_json:
            data = request.get_json()
        else:
            data = request.form
        
        if not data or 'text' not in data:
            return jsonify({"error": "Text parameter required"}), 400
        
        api_key = request.headers.get('X-API-Key') or data.get('api_key')
        
        if api_key:
            db = get_db()
            key_data = db.execute('SELECT * FROM api_keys WHERE key = ? AND is_active = 1', (api_key,)).fetchone()
            
            if not key_data:
                return jsonify({"error": "Invalid or inactive API key"}), 401
            
            db.execute('UPDATE api_keys SET last_used = ? WHERE id = ?', 
                       (datetime.now().isoformat(), key_data['id']))
            db.commit()
        
        text = data['text']
        result = detect_with_patterns(text)
        confidence = calculate_confidence(result)
        
        return jsonify({
            "text": text,
            "is_offensive": result['is_offensive'],
            "offensive_count": result['count'],
            "confidence": confidence,
            "details": {
                "explicit_words": result['found'],
                "hidden_patterns": result['patterns_found'],
                "suspicious_words": result['suspicious_words']
            }
        })
    except Exception as e:
        print(f"❌ API analyze error: {e}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    init_db()
    print("=" * 60)
    print("🛡️  TalkKindly - Complete System (FINAL VERSION)")
    print("=" * 60)
    print("👑 Admin Panel: http://127.0.0.1:5000/admin")
    print("🌐 Main App: http://127.0.0.1:5000")
    print("=" * 60)
    print("✅ All Features Active:")
    print("   • User Authentication & Admin System")
    print("   • Complete Admin Panel with Statistics") 
    print("   • User Management (View/Delete/Disable)")
    print("   • Offensive Words Management")
    print("   • Full History with User Names")
    print("   • Delete Analysis Records (Admin Only)")
    print("   • Activity Logging")
    print("   • API Keys System")
    print("   • Export (PDF/CSV)")
    print("   • File Analysis (CSV, PDF, DOCX, TXT)")
    print("=" * 60)
    app.run(debug=True, port=5000, host='127.0.0.1')